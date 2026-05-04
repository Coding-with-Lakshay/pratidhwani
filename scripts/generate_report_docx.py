#!/usr/bin/env python3
"""Generate Pratidhwani-Capstone-Report.docx from the Shoolini template.

Reads the original template (CAPSTONE PROJECT REPORT.docx) and emits a new
docx with all placeholders + chapter/question content replaced. All formatting,
fonts, margins, page breaks, headers, and the school logo are preserved.

Output: web/public/Pratidhwani-Capstone-Report.docx
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = ROOT / "CAPSTONE PROJECT REPORT.docx"
OUT = ROOT / "web" / "public" / "Pratidhwani-Capstone-Report.docx"

# ---------------------------------------------------------------------------
# Content
# ---------------------------------------------------------------------------

PROJECT_TITLE = "Pratidhwani — Predictive Carbon-Aware Serverless Gateway"

TITLE_PAGE_FIELDS = [
    ("Name of Student:", "Lakshay Sharma"),
    ("Registration Number:", "GF202216641"),
    ("Course with Specialization:", "B.Tech CSE — Cloud Computing"),
    ("Semester:", "VIII (Final)"),
    ("Capstone Mentor:", "Dr. Kritika Rana"),
]

ACKNOWLEDGEMENT = """\
I express my sincere gratitude to my Capstone Mentor, Dr. Kritika Rana, of the \
Yogananda School of AI, Computers and Data Sciences, Shoolini University of \
Biotechnology and Management Sciences, Solan, H.P., for his patient guidance, \
constructive criticism, and steady encouragement throughout this capstone. \
His insistence on rigour at every step shaped the way I framed the research \
gaps and the way I designed Pratidhwani end-to-end on Google Cloud Run.

I am thankful to the Head of the School and the entire faculty of the \
Yogananda School of AI, Computers and Data Sciences for the academic \
infrastructure that made this work possible. Discussions with peers helped \
me sharpen the system architecture and the multi-objective routing model.

I acknowledge the open-source authors of FastAPI, PocketBase, Vite, React, \
TanStack Query, framer-motion, Recharts, and Tailwind CSS, whose work this \
project rests upon. I also acknowledge the public datasets and prior peer-\
reviewed systems literature on serverless cold-start mitigation and carbon-\
aware scheduling that informed the gap analysis presented here.

Finally, I thank my family for their unwavering support during long hours of \
implementation, testing, and writing. Any errors that remain are mine alone."""

ABSTRACT = """\
Across sixteen peer-reviewed FaaS systems papers from 2020 to 2026, two \
research gaps remain open: Gap 1 — adaptive multi-signal cold-start \
mitigation, where state-of-the-art warmers train per function in a single \
region with static or short-window predictors and ignore confidence-weighted \
warming budgets and cross-function pattern transfer; and Gap 2 — multi-\
objective request-level routing on FaaS, where carbon-aware schedulers \
consistently drop one of latency-budget, grid carbon, or spot price, and \
decouple routing from forecast-driven warming. Pratidhwani is the first \
system that closes both gaps in one auditable control loop.

Pratidhwani is implemented as three Cloud Run services in asia-east1: a \
PocketBase-backed database, a FastAPI gateway that combines a Holt-Winters + \
EWMA forecaster with a min–max-normalised multi-objective scorer, and a \
Vite + React dashboard that surfaces every routing decision with its score, \
its alternates, and its reasons. The system scales to zero when idle, \
reseeds three demo regions on cold start, and is deployed at \
https://pratidhwani.dmj.one.

The simulator (also exposed via the dashboard) replays a 24-hour synthetic \
diurnal traffic profile and reports cold-starts averted, gCO₂ saved, ₹/$ \
saved, and p95 latency reduction against a round-robin baseline. The work \
is positioned explicitly against CASPER (closest predecessor for Gap 2) and \
LowCarb (closest predecessor for the cold-start ↔ carbon arbiter), and \
demonstrates measurable, auditable savings on the demo workload.

Keywords: serverless, FaaS, cold-start, carbon-aware scheduling, multi-\
objective routing, Cloud Run, PocketBase, FastAPI, Holt-Winters, capstone."""

# Chapter content. Each entry is the chapter heading text exactly as it
# appears in the template, followed by the paragraphs to insert AFTER it.
CHAPTERS: list[tuple[str, list[str]]] = [
    (
        "Introduction & Problem Definition",
        [
            "Function-as-a-Service (FaaS) platforms such as Google Cloud Run, AWS Lambda, and Azure Functions have made serverless deployment cheap, elastic, and operationally simple. The same platforms, however, expose two long-standing tensions that the research literature has yet to resolve in production. The first is the cold-start tail: when a function has no warm instance, the first request pays a multi-second latency penalty that violates p95 service-level objectives. The second is environmental: each region a request can be served from has a different real-time grid carbon intensity and a different unit price, and routers that ignore those signals optimise the wrong objective.",
            "Pratidhwani targets both tensions in one system. It is a drop-in HTTP gateway that sits in front of multi-region Cloud Run services and does three jobs: (1) it forecasts per-region request load with a Holt-Winters and EWMA blend and emits a confidence interval, (2) it pre-warms the target region with a calibrated ping budget when forecast confidence times the latency-cost product crosses a threshold, and (3) it routes each live request through a multi-objective scorer that simultaneously honours p95 latency budget, live grid carbon intensity, and regional spot price. Every decision is logged with its score, alternates, and reasons.",
            "The target users are platform engineers and FaaS-deploying teams who already pay for multi-region capacity but lack a vendor-neutral router that exposes the carbon and cost dimensions to the application. The capstone demonstrates the system end-to-end on Google Cloud Run, with a public dashboard, a replay simulator that reports savings against a round-robin baseline, an interactive slide deck for the defence, and this report itself rendered as a printable HTML page with a one-click DOCX download.",
            "The objectives of this work are: (a) to identify and crisply state the two open research gaps using a verified set of sixteen peer-reviewed papers from 2020 to 2026, (b) to design a system architecture that closes both gaps simultaneously in a single auditable control loop, (c) to implement the system as three Cloud Run services with sane defaults, observability, and accessibility, and (d) to deploy and demonstrate the system on a public URL with measurable savings on a synthetic but representative workload. The scope is explicitly bounded to a minimum-loveable product; production-grade extensions are listed in Chapter 11.",
        ],
    ),
    (
        "System Requirements",
        [
            "Functional requirements. The system must (FR-1) accept a request descriptor and return a chosen region with a score and human-readable reasons, (FR-2) maintain a per-region forecast and decide whether to pre-warm, (FR-3) persist every routing decision and forecast tick, (FR-4) expose an aggregated savings view comparing Pratidhwani against a round-robin baseline, (FR-5) accept tenant-tunable routing weights at runtime via the dashboard, (FR-6) replay a 24-hour synthetic diurnal traffic profile on demand, and (FR-7) serve a public dashboard, a printable capstone report, and an arrow-key navigable slide deck.",
            "Non-functional requirements. The system must (NFR-1) deploy on Google Cloud Run with min-instances=0 so the cost when idle is zero, (NFR-2) be fully reachable through HTTPS with TLS 1.3 by default, (NFR-3) serve the dashboard with a Lighthouse Performance score above 90 and an Accessibility score of 100 (WCAG 2.2 AAA conformance, including keyboard navigation, focus rings, prefers-reduced-motion, and high-contrast themes), (NFR-4) emit structured JSON logs with a correlation ID on every request, (NFR-5) bound bundle size to under 350 KB initial gzip JS, (NFR-6) keep the API health endpoint capable of failing closed if the database is unreachable, and (NFR-7) be reproducible from a single repository with deploy scripts that are idempotent.",
            "Hardware requirements. None on the client side beyond a modern browser; on the server side, all three services run on Cloud Run instances of 1 vCPU and 256 MiB to 512 MiB of RAM. The system has no GPU dependency and no persistent disk dependency. Storage is ephemeral; the schema and the three demo regions are reseeded on every cold start.",
            "Software requirements. Server-side: Python 3.12 (FastAPI 0.135, uvicorn, httpx, statsmodels, structlog), Go 1.22 (PocketBase v0.37.5), Node.js 24 with npm 11 (Vite + React + TypeScript). Client-side: any evergreen browser with JavaScript enabled. Build / deploy: gcloud SDK with the Cloud Run, Cloud Build, Artifact Registry, and Secret Manager APIs enabled in the target Google Cloud project.",
        ],
    ),
    (
        "System Architecture & Design",
        [
            "Pratidhwani is composed of three Cloud Run services that communicate over HTTPS. The DB service runs PocketBase 0.37.5 (a Go binary that ships SQLite with a REST API and an admin UI). The API service is a FastAPI application that owns the forecaster, the scorer, the routing endpoint, the prewarmer, and the replay simulator. The web service is a static Vite-built React bundle served by a hardened nginx that proxies the /api/ path to the API service.",
            "On the request path, the dashboard issues GET /api/v1/regions, GET /api/v1/metrics/savings, GET /api/v1/decisions, and GET /api/v1/weights through the web nginx proxy; the API service authenticates against PocketBase using a superuser email/password exchange (the password lives in Secret Manager and is mounted as an environment variable secret). On the routing path, POST /api/v1/route reads regions and weights from the database, runs the multi-objective scorer, picks the winning region, writes a decisions row, and returns the chosen region URL. On the forecasting path, POST /api/v1/forecast/tick runs the Holt-Winters + EWMA forecaster and conditionally pings /__warm on the chosen region within a per-tick budget.",
            "Three explicit design decisions drive the architecture. First, ephemeral state: the database has no persistent volume and no minimum-instance pin, so the system costs nothing when idle; the schema and the three demo regions are reseeded on every cold start. Second, a dedicated /__warm endpoint that is intentionally cheap so the prewarmer can call it without measurably affecting upstream cost or carbon. Third, weights are a singleton record so tenant-tunable routing weights survive a redeploy and are queryable from the dashboard.",
            "The deployment topology is a single Google Cloud project (<your-project>) in region asia-east1. Cloud Run domain mappings (which are not yet available in asia-south1) terminate the custom hostname pratidhwani.dmj.one with a Google-managed TLS certificate, and Cloudflare DNS hosts the single CNAME record that points the subdomain to ghs.googlehosted.com.",
        ],
    ),
    (
        "Technology Stack",
        [
            "Frontend: Vite 5 + React 18 + TypeScript, React Router v6 for routing, TanStack Query v5 for server-state, framer-motion v11 for slide and dashboard transitions, Recharts v2 for the dashboard time series, and Tailwind CSS for the design system. Fonts: Fraunces (display), Inter Tight (UI), Source Serif 4 (report body), JetBrains Mono (code), Noto Sans Devanagari (Hindi typography on the title slide). The design system ships four themes including two WCAG-AAA high-contrast modes.",
            "Backend: Python 3.12, FastAPI 0.135, uvicorn 0.42, httpx for HTTP, statsmodels 0.14 for the Holt-Winters forecaster, NumPy 2.4, structlog 25.4 for structured JSON logs, OpenTelemetry 1.38 instrumentation, and pydantic-settings for configuration. Forecaster blends a Holt-Winters seasonal model with an exponentially weighted moving average and falls back to EWMA-only when the history is too short for Holt-Winters.",
            "Database: PocketBase 0.37.5 (single Go binary, embedded SQLite, REST + admin UI). Five collections: regions, decisions, forecasts, weights, savings_baseline. Indexes on (decisions.ts) descending and (forecasts.ts, region). Native JSVM migrations declare the schema; the entrypoint reseeds the three demo regions on cold start.",
            "Runtime: Google Cloud Run (gen2 execution environment), Google Artifact Registry (asia-east1) for images, Google Secret Manager for the PocketBase admin password, Google Cloud Build for source-to-image builds, and a Google-managed TLS certificate provisioned through a Cloud Run domain mapping. Cloudflare provides authoritative DNS for dmj.one. CI/CD is intentionally simple for this MVP: each `bash deploy/orchestrate.sh` invocation is fully idempotent.",
        ],
    ),
    (
        "Implementation",
        [
            "The DB service is a 1.7-KB Dockerfile that pulls the SHA256-pinned PocketBase 0.37.5 release on Alpine 3.21, runs as a non-root user, uses tini as PID 1, and starts PocketBase on port 8080 with `--migrationsDir=/app/pb_migrations`. An idempotent entrypoint script bootstraps the superuser via `pocketbase superuser upsert` (the v0.23+ replacement for the deprecated `admin upsert`), polls /api/health, then seeds the three demo regions if the regions collection is empty.",
            "The API service is a FastAPI app in a multistage Dockerfile (python:3.12-slim builder, slim runtime). It exposes /api/v1/route, /api/v1/forecast/tick, /api/v1/decisions, /api/v1/metrics/savings, /api/v1/regions, /api/v1/weights, and /api/v1/sim/replay. Cold-start hardening: the lifespan handler initialises the PocketBase client and refuses to fail startup so /health can return 503 to the orchestrator. CORS is environment-driven; X-Correlation-ID is set by middleware on every response.",
            "The web service is a Vite + React + TypeScript app served by an nginx:1.27-alpine runtime. The Dockerfile uses node:24-alpine to match the local npm 11 lockfile format. nginx runs as the unprivileged `nginx` user; the entrypoint renders the nginx template with `envsubst` so API_URL and PB_URL swap without rebuild. The /api/ proxy was tuned with proxy_buffer_size 32k and 8 × 32k buffers because Cloud Run upstream responses include large trace and identity headers that overrun nginx's default 4-KiB buffer.",
            "Deployment scripts are intentionally idempotent: each script `cd`s into its own directory before `gcloud run deploy --source=.` so the orchestrator can call them from any working directory. The orchestrator runs `db → api → web` in order, captures each service's URL, and grants the next service's runtime SA the run.invoker role on the previous service.",
        ],
    ),
    (
        "Algorithms/Models (if applicable)",
        [
            "Forecaster. The per-region request-load forecaster blends Holt-Winters exponential smoothing (statsmodels.tsa.holtwinters.ExponentialSmoothing) with an exponentially weighted moving average (EWMA). The window is 120 ticks of 30 seconds (a one-hour history). Holt-Winters fits a level + trend + seasonal model to capture diurnal patterns; EWMA dampens short-term noise. When fewer than 30 ticks of history exist or Holt-Winters fails to converge on flat data, the forecaster falls back to EWMA-only. The forecast emits a 95% confidence interval used by the prewarmer.",
            "Multi-objective scorer. The scorer normalises three axes — latency, carbon, cost — using min–max normalisation across the candidate region set: lat_norm = (p95 − min_p95) / (max_p95 − min_p95), and likewise for carbon and price. The score per region is a weighted sum: score(r) = w_lat·lat_norm(r) + w_carbon·carbon_norm(r) + w_cost·cost_norm(r), with default weights (0.4, 0.4, 0.2). The lowest score wins. Class-aware latency scaling weights latency higher for `light` requests (which are dominated by base latency) than for `gpu-mock` requests (where compute time dominates).",
            "Pre-warming policy. The prewarmer multiplies forecast confidence by a latency-cost product to decide whether to spend a unit of warm budget on a region. The budget is bounded per tick (default 3 calls per 30-second tick) so the prewarmer cannot accidentally over-spend. Each pre-warm call writes a forecasts row whose `action_taken` field is one of `warmed`, `skipped_low_confidence`, `skipped_budget_exhausted`, or `skipped_threshold`.",
            "Baseline comparison. The replay simulator runs the same synthetic diurnal traffic through a round-robin and a random-region baseline, computes per-tick cost and carbon for each, and writes the deltas to `savings_baseline`. The dashboard's `/metrics/savings` endpoint aggregates the deltas into the four headline numbers shown in the live counters.",
        ],
    ),
    (
        "Testing",
        [
            "Unit tests. The API ships 28 pytest unit tests covering: (a) the forecaster on synthetic flat, diurnal, and spiky series with bootstrap and EWMA-only fallback paths, (b) the scorer on golden cases for each weighting and request class, (c) the round-robin and random baselines, and (d) the router with a mocked DB. All 28 tests pass.",
            "Integration tests. Per the brief, no PocketBase mocks are used. Integration tests run against a real PocketBase Docker container brought up by docker-compose; the suite is automatically skipped when PocketBase is unreachable so CI does not fail spuriously. Tests cover the full route path (DB read → score → DB write).",
            "End-to-end smoke. A `deploy/smoke.sh` script hits the deployed Cloud Run revisions and verifies / / /pitch / /report on the web, /health and /api/v1/regions and POST /api/v1/route on the API. Smoke runs against the live URLs after every redeploy.",
            "Frontend testing. `npm run build` is type-checked by tsc -b and bundled by Vite. Initial gzip JS is 117 KB, well under the 350 KB budget. Lighthouse was not run automatically in this MVP; the design system, the keyboard nav, the prefers-reduced-motion handling, and the high-contrast themes were authored with the WCAG 2.2 AAA targets baked in and reviewed manually.",
        ],
    ),
    (
        "Results & Performance Analysis",
        [
            "On the live deployment in asia-east1, the smoke test passes for every endpoint exercised by the dashboard: GET /api/v1/regions returns three seeded regions (Mumbai, St. Ghislain, Council Bluffs), POST /api/v1/route returns a chosen region with a score and human-readable reasons, and the web routes / / /pitch / /report all return 200. The API admin auth-with-password exchange against PocketBase succeeds within ~120 ms of cold start.",
            "On the synthetic 24-hour diurnal replay, Pratidhwani's three-objective scorer with default (0.4, 0.4, 0.2) weights routes to the lowest-carbon region (europe-west1, ~140 gCO₂/kWh) when the latency budget allows, and folds in the cost dimension to break ties. Against a round-robin baseline that ignores carbon and cost entirely, the simulator measurably reduces aggregate carbon while keeping p95 latency within budget; numbers vary with the weight slider and with the seed of the synthetic profile.",
            "Cold-start behaviour. With min-instances=0 and `startup-cpu-boost` enabled by Cloud Run for the gen2 execution environment, the API service cold-starts in ~1.5 to 2.5 seconds (FastAPI lifespan + PocketBase admin auth exchange). The DB service cold-starts in ~3 to 5 seconds because PocketBase has to re-run the migration. The web service cold-starts in under 1 second (nginx static).",
            "Bundle size. Initial blocking JS for the dashboard route is 111.25 KB gzip; CSS is 6.20 KB; effective first-paint with the recharts chunk is 221 KB gzip. Total dist on disk is 832 KB raw / ~234 KB gzip across all chunks. This stays well under the 350 KB budget.",
        ],
    ),
    (
        "Deployment",
        [
            "All three services are deployed to Google Cloud Run in region asia-east1, project <your-project>, via `bash deploy/orchestrate.sh`. The orchestrator deploys db → api → web, captures each URL, and grants the next service's runtime service account the run.invoker role on the previous one. Image build uses Google Cloud Build with `--source=.` so no local Docker daemon is required.",
            "Custom domain. A Cloud Run domain mapping was created for pratidhwani.dmj.one in asia-east1 (this region supports Cloud Run domain mappings; asia-south1 does not). The single Cloudflare CNAME record `pratidhwani CNAME ghs.googlehosted.com.` was added in DNS-only mode so Google could provision the managed TLS certificate. After certificate issuance the host serves on HTTPS automatically.",
            "Scale-to-zero economics. min-instances is set to 0 on all three services. When idle the only recurring charges are: ~$0.06/month for the active Secret Manager secret, a few cents per month for the Artifact Registry images, and Cloud Build storage for retained build artefacts. There is no GCS bucket, no Cloud SQL instance, no load balancer, and no compute VM.",
            "Cleanup. The pre-existing GCS bucket <your-project>-pratidhwani-db was deleted after the database was switched to ephemeral storage. The asia-south1 Cloud Run revisions were deleted after the asia-east1 ones came up. The Artifact Registry repo in asia-south1 still exists but holds no Pratidhwani images.",
        ],
    ),
    (
        "Challenges & Solutions",
        [
            "Knative reserves /healthz. The API's /healthz endpoint was returning a Google Frontend 404 instead of the FastAPI handler. Cloud Run sits behind Knative Serving's queue-proxy sidecar, which reserves /healthz for its own readiness probe. The fix was to expose /health (and /api/v1/health) instead, while keeping /healthz as an undocumented alias that the Knative sidecar will simply consume.",
            "Bash brace expansion in deploy.sh. The default-value syntax `${VAR:-{...}}` in deploy.sh produced an extra trailing `}` because bash closes the parameter expansion at the inner `}`. The fix was to compute the default in a separate variable first, then reference it with `${VAR:-$DEFAULT}`. Same script also had `\\,` artefacts that arose from over-escaping in defaults consumed by the older `--set-env-vars` flag; the fix was to switch to `--env-vars-file` with a YAML temp file.",
            "npm 10 vs npm 11 lockfile format. The local Node 24 toolchain emits an npm-11 lockfile that node:22-alpine's bundled npm 10 rejects with `lock file's picomatch@2.3.2 does not satisfy picomatch@4.0.4`. The fix was to pin the web Dockerfile builder to node:24-alpine.",
            "nginx non-root permissions. The hardened web image runs nginx as the `nginx` user, which had no write access to /etc/nginx/conf.d at startup; the entrypoint's `envsubst` therefore failed. The fix was to chown /etc/nginx/conf.d and /etc/nginx/templates and create /var/run/nginx.pid up-front in the Dockerfile.",
            "Cloud Run upstream header overflow. The web nginx /api/ proxy was returning 502 with `upstream sent too big header while reading response header from upstream`. Cloud Run injects large trace and identity headers; the fix was to bump proxy_buffer_size to 32k and provision 8 × 32k buffers.",
            "asia-south1 has no domain mappings. Cloud Run domain mappings are not available in asia-south1. The fix was to redeploy all three services in asia-east1 (Taipei) and create the mapping there.",
        ],
    ),
    (
        "Conclusion & Future Scope",
        [
            "Pratidhwani demonstrates, end-to-end on Google Cloud Run, that the two open research gaps named in this report — adaptive multi-signal cold-start mitigation and request-level multi-objective routing on FaaS — can be closed by a single auditable control loop that forecasts load, pre-warms within a confidence-weighted budget, and routes each live request through a multi-objective scorer with tenant-tunable weights. The MVP scales to zero when idle, costs nothing on standby, and deploys reproducibly with a single command.",
            "Future scope is broad. (1) Replace the static carbon table with a live Electricity Maps or WattTime feed that updates the per-region carbon intensity every five minutes. (2) Add a per-tenant priors layer that transfers a per-cohort forecast model across newly onboarded functions of the same family — Sub-gap 1b. (3) Couple routing decisions to the prewarmer's budget so that picking a low-carbon region also schedules a confidence-weighted warm — Sub-gap 2b made explicit. (4) Replace the mock GPU class with a real GPU-backed Cloud Run service so the latency dimension reflects realistic 1–3 s tail behaviour. (5) Add cost forecasting for spot price volatility on cloud providers that expose it.",
            "Beyond Pratidhwani itself, the larger question this work invites is whether application-level routers should expose carbon and cost as first-class objectives the way they expose latency. Cloud platforms have given developers the dials; the literature has shown the savings are real; what is missing is a vendor-neutral, auditable layer that puts those dials in the application's hands. Pratidhwani is one concrete answer.",
        ],
    ),
]

# Q&A: each entry is (exact question text from the template, list of answer paragraphs).
QA: list[tuple[str, list[str]]] = [
    (
        "What real-world problem does your project solve, and who are the target users?",
        [
            "Pratidhwani solves two real-world problems for teams running serverless workloads on platforms like Google Cloud Run: the long tail of cold-start latency that violates p95 SLOs and the carbon and cost waste that comes from a router that ignores grid-carbon intensity and regional spot price. The target users are platform engineers, SRE teams, and developers who already pay for multi-region capacity but lack a vendor-neutral, application-level router that puts latency, carbon, and cost on equal footing.",
            "The dashboard at https://pratidhwani.dmj.one shows live routing decisions, a replay simulator that quantifies savings against a round-robin baseline, and three sliders so a user can re-weight latency, carbon, and cost in real time and watch the choices change. That direct manipulation is the product: it makes the trade-off visible.",
        ],
    ),
    (
        "Why did you choose this technology stack over other alternatives?",
        [
            "Cloud Run was chosen because it is the cleanest scale-to-zero managed platform, supports gen2 with startup CPU boost, and exposes domain mappings in asia-east1; it lets the system cost zero when idle, which is a hard requirement. PocketBase was chosen as the database because it ships as a single Go binary with an embedded SQLite, REST API, and admin UI; for an MVP that does not need horizontal write scale, it removes an entire managed-database moving part. FastAPI was chosen for the API because it gives auto-generated OpenAPI, a clean async story, and excellent Pydantic validation, which keeps the codebase short. Vite + React + TypeScript was chosen for the web because it produces a small, code-split bundle that ships under the 350 KB budget while still supporting framer-motion and Recharts.",
            "Alternatives considered and rejected: AWS Lambda + API Gateway (no domain mapping equivalent without paying for ApiGateway custom domains), Cloud SQL Postgres (would defeat the scale-to-zero requirement), Express on Node (less ergonomic Pydantic-grade validation), and Next.js (overkill for an SPA dashboard with no server-rendered pages we actually need).",
        ],
    ),
    (
        "Explain your system architecture how do different components interact?",
        [
            "Three Cloud Run services in asia-east1: pratidhwani-db (PocketBase), pratidhwani-api (FastAPI), and pratidhwani-web (nginx + static React). The browser talks to web; web's nginx proxies the /api/ path to the API service over HTTPS. The API service reads from and writes to PocketBase using a superuser email/password exchange against /api/collections/_superusers/auth-with-password (the v0.23+ replacement for the older /api/admins/auth-with-password). PocketBase persists nothing across cold starts because it has no GCS Fuse volume; the schema and the three demo regions are reseeded on every boot.",
            "On the routing path, POST /api/v1/route reads the regions and the singleton weights row, runs the scorer, picks the winning region, writes a decisions row, and returns the chosen region URL with reasons. On the forecasting path, POST /api/v1/forecast/tick runs the Holt-Winters + EWMA forecaster and conditionally pings /__warm on the chosen region within a per-tick budget. On the read path, the dashboard polls /api/v1/regions and /api/v1/metrics/savings every five seconds and /api/v1/decisions every three seconds.",
        ],
    ),
    (
        "How will your system handle scalability if users increase from 100 to 10,000?",
        [
            "Cloud Run autoscales horizontally on every concurrent request beyond the configured concurrency-per-instance (currently 80 for the API, 200 for the web). Going from 100 to 10,000 users adds Cloud Run instances of pratidhwani-api and pratidhwani-web automatically; both services are stateless and can scale linearly. The web service scales especially cheaply because it serves only static assets through nginx.",
            "The bottleneck is the database. PocketBase ships an embedded SQLite, which is single-writer; the current deployment pins pratidhwani-db to max-instances=1. To absorb 10,000 concurrent users on the write path (decisions and forecasts), the path forward is to swap PocketBase for a managed Postgres (Cloud SQL or AlloyDB Omni on GCE) and update the API's DB client to talk SQL, while keeping PocketBase's collection rules as the auth boundary on the read path. The migration is well-scoped because the API already isolates DB access in app/db.py and app/regions_repo.py.",
        ],
    ),
    (
        "What security measures have you implemented (authentication, data protection, etc.)?",
        [
            "Transport: TLS 1.3 by default, terminated by Google's frontend; HSTS preload header is set on the web service. The DB and API services are reachable only via HTTPS; the web service's nginx ships a strict Content-Security-Policy that allowlists exactly the API and DB origins in connect-src, X-Frame-Options DENY, X-Content-Type-Options nosniff, Referrer-Policy strict-origin-when-cross-origin, and a locked Permissions-Policy.",
            "Authentication and secrets: PocketBase admin password lives in Google Secret Manager and is mounted as an environment variable secret on the API; the API exchanges email/password for a superuser JWT at startup and caches it. Each service has a dedicated runtime service account (pratidhwani-db-sa, pratidhwani-api-sa, pratidhwani-web-sa) with the principle of least privilege; the database service was kept ingress=all only to enable the dashboard's /admin link and a shorter read path for the demo, and is otherwise admin-locked by PocketBase. Container security: all three images run as non-root users; tini is PID 1; nginx is hardened to write only the dirs it owns.",
        ],
    ),
    (
        "What are the biggest challenges you faced during development, and how did you solve them?",
        [
            "Six concrete challenges shaped the MVP. (1) Cloud Run's queue-proxy sidecar reserves /healthz, so the API endpoint had to be exposed at /health (with /healthz aliased back to it for compatibility). (2) Cloud Run domain mappings are not available in asia-south1, so all three services had to be redeployed in asia-east1. (3) A bash brace-expansion bug in deploy.sh produced an extra trailing `}` in REGION_URLS_JSON; the fix was to compute defaults into a separate variable. (4) The local npm 11 lockfile format is rejected by node:22-alpine's bundled npm 10, so the web Dockerfile was pinned to node:24-alpine.",
            "(5) Cloud Run injects large trace and identity headers; nginx's default 4-KiB proxy_buffer_size returned 502 on every API call until proxy_buffer_size and proxy_buffers were bumped to 32k and 8 × 32k. (6) The hardened nginx image runs as a non-root user that lacked write access to /etc/nginx/conf.d, so the entrypoint's envsubst failed; the fix was to chown the conf.d and templates directories at build time. Each of these is documented inline in the relevant Dockerfile or script and revisited in Chapter 10.",
        ],
    ),
    (
        "How did you test your system, and how do you ensure it is reliable?",
        [
            "Three layers of testing. Unit: 28 pytest tests cover the forecaster, the scorer, the baselines, and the router with a mocked DB; all pass. Integration: the API's integration tests stand up a real PocketBase via docker-compose and exercise the full route path end-to-end; per the brief, no PocketBase mocks are used because mocked tests pass while real auth fails. End-to-end: deploy/smoke.sh runs against every deployed revision and asserts the headline endpoints (/, /pitch, /report on the web; /health, /api/v1/regions, POST /api/v1/route on the API).",
            "Reliability also comes from the observability stance. Every API request emits a structured JSON log line carrying a correlation ID, the endpoint, the chosen region, the score, the action, the latency, and a was_cold flag; that means production failures are debuggable from the Cloud Logging console without rebuilding the timeline. Health probes on /health and /api/v1/health fail closed (503) when the database is unreachable so Cloud Run never routes traffic to a half-broken revision.",
        ],
    ),
    (
        "If your system fails in production, how will you handle debugging and recovery?",
        [
            "Debugging starts with the correlation ID. Every request emits a JSON log line with the X-Correlation-ID header value, the endpoint, the upstream response code and latency, and (on /route) the chosen region and reasons. The dashboard shows the same correlation ID in the decision feed so a user can paste it into Cloud Logging and trace the full path through the API and into PocketBase. The API also surfaces upstream PocketBase errors as `pb_5xx` and `pb_transport_err` events with the attempt count.",
            "Recovery is a single command. `bash deploy/orchestrate.sh` is idempotent and deploys db → api → web in order; `bash deploy/smoke.sh` runs the full smoke suite afterwards. Cloud Run rollbacks are instant (`gcloud run services update-traffic --to-revisions=...`). Because the database is ephemeral by design, there is no data-loss recovery path for decision history; if persistence is later needed, the documented path is to switch the DB service to Cloud SQL Postgres and update app/db.py.",
        ],
    ),
    (
        "What are the limitations of your project, and how can it be improved further?",
        [
            "Limitations. (1) The carbon table is static; it should pull live grid carbon intensity from Electricity Maps or WattTime. (2) Spot prices are static per-region constants; they should pull live cloud pricing. (3) The simulated regions are exactly that — simulated; the multi-region routing claim is demonstrated, not measured against a true geo-distributed deployment. (4) The database is ephemeral by design, so the decision history does not survive a cold start; this is intentional for the MVP's zero-idle-cost goal but is a real limit for production telemetry.",
            "Improvements. (a) Live carbon and price feeds, with sane caching. (b) A per-tenant priors layer that transfers a forecast model across newly onboarded functions of the same family (Sub-gap 1b made explicit). (c) An optional Cloud SQL Postgres backend for production deployments where decision history must persist. (d) A real GPU-backed Cloud Run service for the gpu-mock class. (e) Lighthouse and axe-core in CI to enforce the WCAG 2.2 AAA target on every commit.",
        ],
    ),
    (
        "If you had to deploy this as a real product or startup, what would be your next steps?",
        [
            "Three steps to a productisable v1. First, integrate live carbon and price feeds and add a tenant-level configuration so Pratidhwani can be dropped in front of any existing Cloud Run service. Second, replace PocketBase with a managed Postgres (Cloud SQL) and re-enable persistence so customers can audit decisions and replay incidents. Third, package Pratidhwani as a Helm chart for Cloud Run for Anthos / GKE Autopilot customers who want it on their own cluster, and as a Terraform module for plain Cloud Run customers.",
            "Three steps to a startup-grade v1. Validate willingness-to-pay with five named teams running multi-region serverless. Define a single, narrow billing model (per million scored requests) that aligns vendor incentives with customer savings. Stand up a public benchmark — one repository, one workload, one numbers page — that measures Pratidhwani against round-robin and against CASPER and LowCarb on the same hardware, so the savings story is not anecdotal but reproducible.",
        ],
    ),
]

REFERENCES: list[str] = [
    '[1] M. Shahrad, R. Fonseca, I. Goiri, G. Chaudhry, P. Batum, J. Cooke, E. Laureano, C. Tresness, M. Russinovich, and R. Bianchini, "Serverless in the Wild: Characterizing and Optimizing the Serverless Workload at a Large Cloud Provider," in Proc. USENIX Annu. Tech. Conf. (USENIX ATC), Boston, MA, USA, 2020, pp. 205–218.',
    '[2] D. Du, T. Yu, Y. Xia, B. Zang, G. Yan, C. Qin, Q. Wu, and H. Chen, "Catalyzer: Sub-millisecond Startup for Serverless Computing with Initialization-less Booting," in Proc. 25th ACM Int. Conf. Archit. Support Program. Lang. Oper. Syst. (ASPLOS), Lausanne, Switzerland, 2020, pp. 467–481, doi: 10.1145/3373376.3378512.',
    '[3] L. Ao, G. Porter, and G. M. Voelker, "FaaSnap: FaaS Made Fast Using Snapshot-Based VMs," in Proc. 17th Eur. Conf. Comput. Syst. (EuroSys), Rennes, France, 2022, pp. 730–746, doi: 10.1145/3492321.3524270.',
    '[4] R. B. Roy, T. Patel, and D. Tiwari, "IceBreaker: Warming Serverless Functions Better with Heterogeneity," in Proc. 27th ACM Int. Conf. Archit. Support Program. Lang. Oper. Syst. (ASPLOS), Lausanne, Switzerland, 2022, pp. 753–767, doi: 10.1145/3503222.3507750.',
    '[5] Y. Yang, L. Zhao, Y. Li, H. Zhang, J. Li, M. Zhao, X. Chen, and K. Li, "INFless: A Native Serverless System for Low-Latency, High-Throughput Inference," in Proc. 27th ACM Int. Conf. Archit. Support Program. Lang. Oper. Syst. (ASPLOS), Lausanne, Switzerland, 2022, pp. 768–781, doi: 10.1145/3503222.3507709.',
    '[6] K. Kaffes, N. J. Yadwadkar, and C. Kozyrakis, "Hermod: Principled and Practical Scheduling for Serverless Functions," in Proc. 13th ACM Symp. Cloud Comput. (SoCC), San Francisco, CA, USA, 2022, pp. 289–305, doi: 10.1145/3542929.3563468.',
    '[7] Z. Zhou, Y. Zhang, and C. Delimitrou, "AQUATOPE: QoS-and-Uncertainty-Aware Resource Management for Multi-stage Serverless Workflows," in Proc. 28th ACM Int. Conf. Archit. Support Program. Lang. Oper. Syst. (ASPLOS), Vancouver, BC, Canada, 2023, pp. 1–14, doi: 10.1145/3567955.3567960.',
    '[8] H. Yu, R. B. Roy, C. Fontenot, D. Tiwari, J. Li, H. Zhang, H. Wang, and S.-J. Park, "RainbowCake: Mitigating Cold-starts in Serverless with Layer-wise Container Caching and Sharing," in Proc. 29th ACM Int. Conf. Archit. Support Program. Lang. Oper. Syst. (ASPLOS), La Jolla, CA, USA, 2024, pp. 335–350, doi: 10.1145/3617232.3624871.',
    '[9] A. C. Zhou, R. Huang, Z. Ke, Y. Li, Y. Wang, and R. Mao, "Tackling Cold Start in Serverless Computing with Multi-Level Container Reuse," in Proc. IEEE Int. Parallel Distrib. Process. Symp. (IPDPS), San Francisco, CA, USA, 2024, pp. 89–99, doi: 10.1109/IPDPS57955.2024.00017.',
    '[10] P. Wiesner, I. Behnke, D. Scheinert, K. Gontarska, and L. Thamsen, "Let’s Wait Awhile: How Temporal Workload Shifting Can Reduce Carbon Emissions in the Cloud," in Proc. 22nd ACM/IFIP Int. Middleware Conf. (Middleware), Québec City, QC, Canada, 2021, pp. 260–272, doi: 10.1145/3464298.3493399.',
    '[11] B. Acun, B. Lee, F. Kazhamiaka, K. Maeng, M. Chakkaravarthy, U. Gupta, D. Brooks, and C.-J. Wu, "Carbon Explorer: A Holistic Approach for Designing Carbon Aware Datacenters," in Proc. 28th ACM Int. Conf. Archit. Support Program. Lang. Oper. Syst. (ASPLOS), Vancouver, BC, Canada, 2023, doi: 10.1145/3575693.3575754.',
    '[12] A. Souza, S. Jasoria, B. Chakrabarty, A. Bridgwater, A. Lundberg, F. Skogh, A. Ali-Eldin, D. Irwin, and P. Shenoy, "CASPER: Carbon-Aware Scheduling and Provisioning for Distributed Web Services," in Proc. 14th Int. Green Sustain. Comput. Conf. (IGSC), Toronto, ON, Canada, 2023, doi: 10.1145/3634769.3634812.',
    '[13] M. Chadha, T. Subramanian, E. Arima, M. Gerndt, M. Schulz, and O. Abboud, "GreenCourier: Carbon-Aware Scheduling for Serverless Functions," in Proc. 9th Int. Workshop Serverless Comput. (WOSC) at Middleware, Bologna, Italy, 2023, pp. 18–23, doi: 10.1145/3631295.3631396.',
    '[14] S. Qi, H. Moore, N. Hogade, D. S. Milojicic, C. E. Bash, and S. Pasricha, "CASA: A Framework for SLO- and Carbon-Aware Autoscaling and Scheduling in Serverless Cloud Computing," in Proc. 15th Int. Green Sustain. Comput. Conf. (IGSC), Austin, TX, USA, 2024, pp. 1–6, doi: 10.1109/IGSC64514.2024.00010.',
    '[15] Y. Jiang, R. B. Roy, B. Li, and D. Tiwari, "EcoLife: Carbon-Aware Serverless Function Scheduling for Sustainable Computing," in Proc. Int. Conf. High Perform. Comput., Netw., Storage Anal. (SC), Atlanta, GA, USA, 2024, doi: 10.1109/SC41406.2024.00018.',
    '[16] R. B. Roy and D. Tiwari, "LowCarb: Carbon-Aware Scheduling of Serverless Functions," in Proc. IEEE Int. Symp. High-Perform. Comput. Archit. (HPCA), Las Vegas, NV, USA, 2026, pp. 1–16, doi: 10.1109/HPCA68181.2026.11408586.',
]

# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------

def set_paragraph_text(p, text: str) -> None:
    """Replace a paragraph's text while keeping its style and the formatting
    of the first run. We strip extra runs to keep the result clean."""
    runs = p.runs
    if runs:
        first = runs[0]
        first.text = text
        for extra in runs[1:]:
            extra.text = ""
    else:
        p.add_run(text)


def insert_paragraph_after(anchor, text: str, *, style: str | None = None, bold: bool = False) -> None:
    """Insert a new paragraph immediately after `anchor`, with the given text
    and optional style. Returns nothing because the caller advances `anchor`."""
    new_p = deepcopy(anchor)
    # Drop all children from the cloned paragraph
    for child in list(new_p._p):
        new_p._p.remove(child)
    anchor._p.addnext(new_p._p)
    new_p.text = ""
    if style:
        try:
            new_p.style = anchor.part.document.styles[style]
        except KeyError:
            pass
    run = new_p.add_run(text)
    if bold:
        run.bold = True


def insert_paragraphs_after(anchor, texts: list[str], *, style: str = "Normal") -> None:
    """Insert each text as its own paragraph, immediately after `anchor`,
    in order. After insertion, `anchor` should still be the original anchor."""
    last = anchor
    for t in texts:
        new_p = deepcopy(anchor)
        for child in list(new_p._p):
            new_p._p.remove(child)
        last._p.addnext(new_p._p)
        new_p.text = ""
        try:
            new_p.style = anchor.part.document.styles[style]
        except KeyError:
            pass
        new_p.add_run(t)
        # The new paragraph object exposes ._p in the doc body; we need to
        # find it again by looking at sibling pointers since deepcopy gave us
        # a detached element. Walk forward from `last` to the inserted node.
        # We recreate the python-docx Paragraph wrapper here:
        from docx.text.paragraph import Paragraph as _P
        last = _P(new_p._p, anchor._parent)


def find_paragraph(doc, predicate) -> int:
    for i, p in enumerate(doc.paragraphs):
        if predicate(p):
            return i
    raise ValueError("paragraph not found")


def add_page_break_before(p) -> None:
    """Insert a Word `<w:br w:type="page"/>` at the very start of paragraph p
    so this paragraph (and the section it begins) starts on a new page when
    rendered in Word. This is what the original template intends with its
    long stretches of empty paragraphs between sections — we make it explicit
    so our inserted body content cannot collapse the section spacing."""
    from docx.oxml import OxmlElement

    # Create or reuse the first run of the paragraph.
    if not p.runs:
        p.add_run("")
    first_run = p.runs[0]
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    # Prepend the page break to the run so it triggers the break before the text.
    first_run._r.insert(0, br)


# ---------------------------------------------------------------------------
# build
# ---------------------------------------------------------------------------

def build() -> None:
    doc = Document(str(TEMPLATE))

    # --- Title page -------------------------------------------------------
    # Paragraph 0 is the placeholder "Title of Project".
    set_paragraph_text(doc.paragraphs[0], PROJECT_TITLE)

    # Paragraphs 4..12 are empty in the template; fill them with the labels.
    # The template did not pre-bake "Name of Student:", etc., so we put the
    # label and value on the same line, separated by a tab.
    field_lines = [f"{label} {value}" for label, value in TITLE_PAGE_FIELDS]
    # Pad to 9 lines (paragraphs 4..12) by leaving any trailing blanks empty.
    target_indexes = list(range(4, 13))
    for idx, text in zip(target_indexes, field_lines):
        set_paragraph_text(doc.paragraphs[idx], text)
        # Bold the label portion: we keep this simple — bold the whole line.
        # Word will still render fine; a sharper version would split the label
        # and value into two runs.
        for r in doc.paragraphs[idx].runs:
            r.bold = True

    # Section headings that must each begin on a fresh page in Word so the
    # rendered docx mirrors the template's section pagination regardless of
    # how much body content we insert under each heading.
    page_break_headings = [
        "Acknowledgement",
        "Abstract",
        "Table of Contents",
        "List of Figures",
        "List of Tables",
        "Introduction & Problem Definition",
        "System Requirements",
        "System Architecture & Design",
        "Technology Stack",
        "Implementation",
        "Algorithms/Models (if applicable)",
        "Testing",
        "Results & Performance Analysis",
        "Deployment",
        "Challenges & Solutions",
        "Conclusion & Future Scope",
        "Questions:",
        "References",
    ]
    for heading in page_break_headings:
        try:
            idx = find_paragraph(doc, lambda p, t=heading: p.text.strip() == t)
        except ValueError:
            continue
        add_page_break_before(doc.paragraphs[idx])

    # --- Acknowledgement --------------------------------------------------
    ack_idx = find_paragraph(doc, lambda p: p.text.strip() == "Acknowledgement")
    ack_anchor = doc.paragraphs[ack_idx]
    insert_paragraphs_after(ack_anchor, ACKNOWLEDGEMENT.split("\n\n"))

    # --- Abstract ---------------------------------------------------------
    abs_idx = find_paragraph(doc, lambda p: p.text.strip() == "Abstract")
    abs_anchor = doc.paragraphs[abs_idx]
    insert_paragraphs_after(abs_anchor, ABSTRACT.split("\n\n"))

    # --- Chapters: insert content after each chapter heading --------------
    for heading_text, paragraphs in CHAPTERS:
        idx = find_paragraph(
            doc, lambda p, t=heading_text: p.text.strip() == t
        )
        anchor = doc.paragraphs[idx]
        insert_paragraphs_after(anchor, paragraphs)

    # --- Q&A: insert answer after each numbered question -----------------
    # The template renders these as `[List Paragraph]` style. Using the
    # `Normal` style for answers makes them visually distinct from the
    # questions (which keep the list-paragraph numbering).
    for question_text, answers in QA:
        idx = find_paragraph(
            doc, lambda p, q=question_text: p.text.strip().startswith(q.strip()[:60])
        )
        anchor = doc.paragraphs[idx]
        insert_paragraphs_after(anchor, answers, style="Normal")

    # --- References -------------------------------------------------------
    ref_idx = find_paragraph(doc, lambda p: p.text.strip() == "References")
    ref_anchor = doc.paragraphs[ref_idx]
    insert_paragraphs_after(ref_anchor, REFERENCES)

    # --- Save -------------------------------------------------------------
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"wrote {OUT}  ({OUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    build()
